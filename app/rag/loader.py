from pathlib import Path
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter


async def load_pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    return text


async def load_docx(path: Path) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    return text


async def load_markdown(path: Path) -> str:
    return path.read_text(encoding="utf-8")


async def load_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8")


LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".md": load_markdown,
    ".txt": load_txt,
}


async def load_documents(directory: str) -> list[Document]:
    """加载目录中所有支持的文档，返回 LangChain Document 列表"""
    docs = []
    dir_path = Path(directory)

    for file_path in dir_path.iterdir():
        if file_path.is_file():
            ext = file_path.suffix.lower()
            loader = LOADERS.get(ext)
            if loader:
                try:
                    text = await loader(file_path)
                    docs.append(Document(
                        page_content=text,
                        metadata={"source": file_path.name, "type": ext}
                    ))
                except Exception as e:
                    print(f"Warning: Failed to load {file_path.name}: {e}")

    # Split into chunks
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?"],
    )
    chunks = splitter.split_documents(docs)
    return chunks
